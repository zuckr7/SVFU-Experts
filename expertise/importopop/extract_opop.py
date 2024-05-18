from docx import Document
import re

def extract_opop(docx_file):
    doc = Document(docx_file)
    opop = {}
    last__num = 0

    def remove_delimiters(text):
        return re.sub(r'\s+', '', text)

    def find_word_in_table(word):
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cell_text = row.cells[0].text.strip()
                cell_text = remove_delimiters(cell_text)
                if word in cell_text:
                    return i
        return None


    def extract_program_name(doc):
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                program_name = table.cell(find_word_in_table("Код"), len(table.columns) - 1).text.strip()
                return program_name.strip(",;.:").strip()
        return None

    def extract_edu_lvl(doc):
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                edu_lvl = table.cell(find_word_in_table("Уровень"), len(table.columns) - 1).text.strip()
                return edu_lvl.strip(",;.:").strip()
        return None


    def extract_UK(doc):
        uk_dict = {}
        n=0
        uk_number=0
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                uk = table.cell(find_word_in_table("Требованиякрезультатам"), len(table.columns) - 1).text.strip()
                uk_list = uk.split("\n")
                for item in uk_list:
                    if 'УК-' in item:
                        uk_number = item.split('УК-')[1].split(')')[0]
                    
                        uk_text = item.replace(f'УК-{uk_number})', '').strip()
                        uk_dict[f'УК-{uk_number}'] = {'name':uk_text.strip(",;.:()").strip()
                                                        ,'INDIC':{}
                                                        ,'ZUV': {'Z': [], 'U': [], 'V': []}
                                                    }
                        for index,table in enumerate(doc.tables):
                            if len(table.rows) > 0 and len(table.columns) > 0 and index==1:
                                uk1 = table.cell(int(uk_number), len(table.columns) - 2).text.strip()
                                uk_list1 = uk1.split("\n")
                                for j,item2 in enumerate(uk_list1):
                                    uk_dict[f'УК-{uk_number}']["INDIC"][f'УК-{uk_number}.{j+1}']= item2.split(" ",1)[-1]
                                zuv = table.cell(int(uk_number), len(table.columns) - 1).text.strip()
                                zuv_list = zuv.split("\n")
                                for k,zuvs in enumerate(zuv_list):
                                    if("Знать:" in zuvs):
                                        current = 'Z'
                                        in_know = True
                                        continue
                                    elif("Уметь:" in zuvs):
                                        current = 'U'
                                        in_know = True
                                        continue
                                    if("Владеть:" in zuvs):
                                        current = 'V'
                                        in_know = True
                                        continue
                                    
                                    elif in_know:
                                        uk_dict[f'УК-{uk_number}']['ZUV'][current].append(zuvs)
                                    elif(":" in zuvs):
                                        in_know=False
                                        break
            return uk_dict, uk_number
        return None



    def extract_OPK(doc, num):
        opk_dict = {}
        opk_number = 0
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                opk = table.cell(find_word_in_table("Требованиякрезультатамосвоенияпрограммы"), len(table.columns) - 1).text.strip()
                opk_list = opk.split("\n")
                for item in opk_list:
                    if '(ОПК-' in item:
                        opk_number = item.split('(ОПК-')[1].split(')')[0]
                        opk_text = item.replace(f'(ОПК-{opk_number})', '').strip()
                        opk_dict[f'ОПК-{opk_number}'] = {"name":opk_text.strip(",;.:").strip()
                                                        ,"INDIC":{}
                                                        ,'ZUV': {'Z': [], 'U': [], 'V': []}
                                                        }
                        for index,table in enumerate(doc.tables):
                            if len(table.rows) > 0 and len(table.columns) > 0 and index==1:
                                opk1 = table.cell(int(opk_number)+int(num), len(table.columns) - 2).text.strip()
                                opk_list1 = opk1.split("\n")
                                for j,item2 in enumerate(opk_list1):
                                    opk_dict[f'ОПК-{opk_number}']["INDIC"][f'ОПК-{opk_number}.{j+1}']= item2.split(" ",1)[-1]
                                zuv = table.cell(int(opk_number)+int(num), len(table.columns) - 1).text.strip()
                                zuv_list = zuv.split("\n")
                                for k,zuvs in enumerate(zuv_list):
                                    if("Знать:" in zuvs):
                                        current = 'Z'
                                        in_know = True
                                        continue
                                    elif("Уметь:" in zuvs):
                                        current = 'U'
                                        in_know = True
                                        continue
                                    if("Владеть:" in zuvs):
                                        current = 'V'
                                        in_know = True
                                        continue
                                    
                                    elif in_know:
                                        opk_dict[f'ОПК-{opk_number}']['ZUV'][current].append(zuvs)
                                    elif(":" in zuvs):
                                        in_know=False
                                        break
                                    
            return opk_dict,int(num)+int(opk_number)
        return None

    def extract_PK(doc , num):
        pk_dict = {}
        pk_number = 0
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                
                pk = table.cell(find_word_in_table("Требованиякрезультатамосвоенияпрограммы"), len(table.columns) - 1).text.strip()
                pk_list = pk.split("\n")
                for item in pk_list:
                    if '(ПК-' in item:
                        pk_number = item.split('(ПК-')[1].split(')')[0]
                        pk_text = item.replace(f'(ПК-{pk_number})', '').strip()
                        pk_dict[f'ПК-{pk_number}'] = {"name":pk_text.strip(",;.:").strip()
                                                    ,"INDIC":{}
                                                    ,'ZUV': {'Z': [], 'U': [], 'V': []}
                                                    }
                        for index,table in enumerate(doc.tables):
                            if len(table.rows) > 0 and len(table.columns) > 0 and index==1:
                                pk1 = table.cell(int(pk_number)+int(num), len(table.columns) - 2).text.strip()
                                pk_list1 = pk1.split("\n")
                                for j,item2 in enumerate(pk_list1):
                                    pk_dict[f'ПК-{pk_number}']["INDIC"][f'ПК-{pk_number}.{j+1}']= item2.split(" ",1)[-1]
                                zuv = table.cell(int(pk_number)+int(num), len(table.columns) - 1).text.strip()
                                zuv_list = zuv.split("\n")
                                for k,zuvs in enumerate(zuv_list):
                                    if("Знать:" in zuvs):
                                        current = 'Z'
                                        in_know = True
                                        continue
                                    elif("Уметь:" in zuvs):
                                        current = 'U'
                                        in_know = True
                                        continue
                                    if("Владеть:" in zuvs):
                                        current = 'V'
                                        in_know = True
                                        continue
                                    
                                    elif in_know:
                                        pk_dict[f'ПК-{pk_number}']['ZUV'][current].append(zuvs)
                                    elif(":" in zuvs):
                                        in_know=False
                                        break
            
            
            return dict(sorted(pk_dict.items()))
        return None


    def extract_disciplines(doc):
        dis_dict = {}
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                dis = table.cell(find_word_in_table("Дисциплины"), len(table.columns) - 1).text.strip()
                dis_list = dis.split("\n")
                for item in dis_list:
                    if "Дисциплины " not in item:
                        if len(item.split(" ", 1)) > 1:
                            
                            dis_code, dis_name = item.replace('\t', " ").split(' ', 1)
                            dis_dict[dis_code] = dis_name.strip(",;.:").strip()
                        
                        
            return dis_dict
        return None

    def extract_practice(doc):
        practice_dict = {}
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) > 0:
                practice = table.cell(find_word_in_table("Практик"), len(table.columns) - 1).text.strip()
                practice_list = practice.split("\n")
                for item in practice_list:
                    if "Учебная практика (концентр.):" not in item and "Производственная практика (концентр.):" not in item:
                        if len(item.split(" ", 1)) > 1:
                            practice_code, practice_name = item.replace('\t', " ").split(' ', 1)
                            practice_dict[practice_code] = practice_name.strip(",;.:").strip()
                        
                        
            return practice_dict
        return None


    def extract_workstandards(doc):
        workstandard_dict = {}
        current_workstandard = None
        current_workfunction_index = -1
        collecting_functions = False
        current_section = ""
        in_know = False
        current_workfunction_name = ""
        fun =0
        
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()


            if "Профессиональный стандарт «" in text:
                name = text.split("«")[1].split("»")[0].strip()
                current_workstandard = name
                workstandard_dict[current_workstandard] = {'workfunctions': {}}
                current_workfunction_index = -1
                
            
            
            elif "уровень" in text:

                current_workfunction_index += 1
                parts = text.split("–")
                workfunction_name = parts[0].strip()
                current_workfunction_name = workfunction_name

                workfunction_level = parts[1].split()[0].strip()
                workstandard_dict[current_workstandard]['workfunctions'][current_workfunction_index] = {
                    'name': workfunction_name,
                    'level': workfunction_level,
                    'functions': {},
                    'ZUV': {'Z': [], 'U': [], 'V': []}
                }

                for j in range(i,len(doc.paragraphs)):
                    if "Трудовые функции:" in doc.paragraphs[j].text.strip():
                        last = False
                        doc.paragraphs[j].text = "zxcursed"
                        fun = j
                        for k in range(j+1,len(doc.paragraphs)):
                            if (doc.paragraphs[k].text.strip() == '') or "Выпускник должен знать" in doc.paragraphs[k].text.strip():
                                break
                            else:
                                functions_dict = workstandard_dict[current_workstandard]['workfunctions'][current_workfunction_index]['functions']
                                function_list = [func.strip() for func in doc.paragraphs[k].text.strip().split('\n') if func.strip()]
                                current_function_index = len(functions_dict)
                                for idx, func in enumerate(function_list):
                                    key = f"{current_workfunction_index}.{current_function_index + idx}"
                                    functions_dict[key] = func

                        
                        break
                for j in range(i,len(doc.paragraphs)):
                    if "Выпускник должен знать" in doc.paragraphs[j].text.strip():
                        current_section = "Z"
                        doc.paragraphs[j].text = "Z"
                        last = False
                        in_know = True
                        continue

                    if "Выпускник должен уметь" in doc.paragraphs[j].text.strip():
                        current_section = "U"
                        doc.paragraphs[j].text = "U"
                        last = False
                        in_know = True
                        continue

                    if "Выпускник должен владеть" in doc.paragraphs[j].text.strip():
                        current_section = "V"
                        doc.paragraphs[j].text = "V"
                        last = True
                        in_know = True
                        
                        continue
                    
                    elif in_know:
                        if(doc.paragraphs[j].text.strip() and last==False):
                            workstandard_dict[current_workstandard]['workfunctions'][current_workfunction_index]['ZUV'][current_section].append(doc.paragraphs[j].text.strip())
                        elif(last):
                            if(doc.paragraphs[j].text.strip()):
                                
                                workstandard_dict[current_workstandard]['workfunctions'][current_workfunction_index]['ZUV'][current_section].append(doc.paragraphs[j].text.strip())
                        
                            else:
                                in_know = False
                                last=False
                                break
                        elif(in_know == True):
                            in_know = False
                    
        return workstandard_dict

    extracted_program_name = extract_program_name(doc).split(" ",1)
    opop['code'] = extracted_program_name[0]
    opop['program_name'] = extracted_program_name[1]

    opop['edu_lvl'] = extract_edu_lvl(doc)
    
    opop['UK'] = {}
    opop['UK'], last_num = extract_UK(doc)

    opop['OPK'] = {}
    opop['OPK'], last_num = extract_OPK(doc,last_num)

    opop['PK'] = {}
    opop['PK'] = extract_PK(doc,last_num)

    opop['disciplines'] = {}
    opop['disciplines'] = extract_disciplines(doc)

    opop['practice'] = {}
    opop['practice'] = extract_practice(doc)


    opop['workstandards']=extract_workstandards(doc)
    return opop


if __name__ == '__main__':
    docx_file = "opop.docx"
    opop = extract_opop(docx_file)









